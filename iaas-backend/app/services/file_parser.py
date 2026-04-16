from io import BytesIO
from zipfile import ZipFile
import xml.etree.ElementTree as ET

import fitz
from docx import Document


def extract_text_from_pdf(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text("text") or "")
        return "\n".join(text_parts).strip()
    finally:
        doc.close()


def extract_text_from_docx(file_bytes):
    document = Document(BytesIO(file_bytes))
    text_parts = []

    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            text_parts.append(text)

    # Many JD templates place core content in tables, so include those cells as well.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = "\n".join(
                    (paragraph.text or "").strip()
                    for paragraph in cell.paragraphs
                    if (paragraph.text or "").strip()
                ).strip()
                if cell_text:
                    text_parts.append(cell_text)

    # Fallback for DOCX layouts where python-docx misses body text
    # (for example, certain templates/content controls/textbox-heavy files).
    if not text_parts:
        text_parts.extend(_extract_text_from_docx_xml(file_bytes))

    unique_parts = list(dict.fromkeys(part for part in text_parts if part))
    return "\n".join(unique_parts).strip()


def _extract_text_from_docx_xml(file_bytes):
    xml_text_parts = []

    with ZipFile(BytesIO(file_bytes)) as docx_zip:
        xml_candidates = [
            name for name in docx_zip.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ]

        for xml_name in xml_candidates:
            try:
                xml_bytes = docx_zip.read(xml_name)
                root = ET.fromstring(xml_bytes)
            except Exception:
                continue

            for element in root.iter():
                # WordprocessingML text node
                if element.tag.endswith("}t"):
                    text = (element.text or "").strip()
                    if text:
                        xml_text_parts.append(text)

    return xml_text_parts